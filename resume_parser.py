"""

resume_parser.py

"""


import io
import re
import csv

import PyPDF2
import docx


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = []
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    for page in reader.pages:
        page_text = page.extract_text() or ""
        text.append(page_text)
    return "\n".join(text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_text(filename: str, file_bytes: bytes) -> str:
    
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}")


EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d{1,3}[\s.-]?)?(\(?\d{3,4}\)?[\s.-]?)\d{3}[\s.-]?\d{3,4}")


def guess_candidate_name(text: str, filename: str) -> str:
    
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if 2 <= len(line.split()) <= 4 and not any(c.isdigit() for c in line) and "@" not in line:
            return line.title()
        break  

    
    name = re.sub(r"\.(pdf|docx)$", "", filename, flags=re.IGNORECASE)
    name = re.sub(r"[_\-]+", " ", name)
    return name.strip().title()


def extract_contact_info(text: str, filename: str) -> dict:
    
    email_match = EMAIL_RE.search(text)
    phone_match = PHONE_RE.search(text)
    return {
        "name": guess_candidate_name(text, filename),
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0).strip() if phone_match else "",
    }



